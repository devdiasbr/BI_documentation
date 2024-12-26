"""
Automação para Geração de Documentação de Projetos Power BI
-----------------------------------------------------------

Este código automatiza a documentação de relatórios Power BI a partir de um arquivo `.pbit` 
convertido em `.zip`. Ele extrai informações de páginas, tabelas, colunas, medidas, fontes, 
e relacionamentos para gerar uma documentação detalhada em Word.
"""

import json
import os 
from os import path, rename
import zipfile
from datetime import datetime
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import StringIO
import time
import logging
from typing import Dict, List, Optional, Union, Any
from pathlib import Path

# Configuração do logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('power_bi_doc.log'),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)

# Tipos customizados para melhor type hinting
JsonDict = Dict[str, Any]
VisualConfig = Dict[str, Any]

# Função para verificar e renomear arquivos
def verificar_ou_renomear_arquivo(arquivo_pbit: str, arquivo_zip: str) -> bool:
    """
    Verifica a existência do arquivo zip e renomeia o arquivo .pbit se necessário.
    
    Args:
        arquivo_pbit (str): Caminho do arquivo .pbit
        arquivo_zip (str): Caminho do arquivo .zip desejado
        
    Returns:
        bool: True se o processo foi bem sucedido, False caso contrário
        
    Raises:
        FileNotFoundError: Se o arquivo .pbit não existir
        PermissionError: Se não houver permissão para renomear o arquivo
    """
    try:
        if os.path.exists(arquivo_zip):
            logger.info(f"Arquivo zip já existe em: {arquivo_zip}")
            return True
            
        if not os.path.exists(arquivo_pbit):
            raise FileNotFoundError(f"Arquivo .pbit não encontrado: {arquivo_pbit}")
            
        os.rename(arquivo_pbit, arquivo_zip)
        logger.info(f"Arquivo renomeado com sucesso: {arquivo_pbit} -> {arquivo_zip}")
        return True
        
    except PermissionError as e:
        logger.error(f"Erro de permissão ao renomear arquivo: {e}")
        raise
    except Exception as e:
        logger.error(f"Erro inesperado ao processar arquivo: {e}")
        return False

# Função para extrair arquivos do ZIP
def extrair_arquivos_zip(arquivo_zip: str, caminho_BI: str, arquivos_para_extrair: List[str]) -> bool:
    """
    Extrai arquivos específicos de um arquivo ZIP.
    
    Args:
        arquivo_zip (str): Caminho do arquivo ZIP
        caminho_BI (str): Diretório onde os arquivos serão extraídos
        arquivos_para_extrair (List[str]): Lista de arquivos a serem extraídos
        
    Returns:
        bool: True se todos os arquivos foram extraídos com sucesso, False caso contrário
        
    Raises:
        FileNotFoundError: Se o arquivo ZIP não existir
        zipfile.BadZipFile: Se o arquivo ZIP estiver corrompido
    """
    try:
        if not os.path.exists(arquivo_zip):
            raise FileNotFoundError(f"Arquivo ZIP não encontrado: {arquivo_zip}")
            
        with zipfile.ZipFile(arquivo_zip, 'r') as zip_ref:
            # Verifica se todos os arquivos existem no ZIP
            zip_files = zip_ref.namelist()
            for arquivo in arquivos_para_extrair:
                if arquivo not in zip_files:
                    raise KeyError(f"Arquivo não encontrado no ZIP: {arquivo}")
            
            # Extrai os arquivos
            for arquivo in arquivos_para_extrair:
                zip_ref.extract(arquivo, caminho_BI)
                logger.info(f"Arquivo extraído com sucesso: {arquivo}")
                
        logger.info(f"Todos os arquivos foram extraídos com sucesso para: {caminho_BI}")
        return True
        
    except zipfile.BadZipFile as e:
        logger.error(f"Arquivo ZIP corrompido: {e}")
        return False
    except KeyError as e:
        logger.error(f"Arquivo não encontrado no ZIP: {e}")
        return False
    except Exception as e:
        logger.error(f"Erro inesperado ao extrair arquivos: {e}")
        return False

# Função para carregar os dados do JSON
def carregar_dados_json(arquivo: str, encoding: str = 'utf-16-le') -> JsonDict:
    """
    Carrega dados de um arquivo JSON com tratamento de erros.
    
    Args:
        arquivo (str): Caminho do arquivo JSON a ser carregado
        encoding (str, optional): Codificação do arquivo. Defaults to 'utf-16-le'
        
    Returns:
        JsonDict: Dicionário com os dados do JSON ou dicionário vazio em caso de erro
        
    Raises:
        FileNotFoundError: Se o arquivo não existir
        json.JSONDecodeError: Se o arquivo não for um JSON válido
    """
    try:
        if not os.path.exists(arquivo):
            raise FileNotFoundError(f"Arquivo não encontrado: {arquivo}")
            
        with open(arquivo, 'r', encoding=encoding) as f:
            dados = json.load(f)
            logger.info(f"Arquivo JSON carregado com sucesso: {arquivo}")
            return dados
            
    except FileNotFoundError as e:
        logger.error(f"Arquivo não encontrado: {e}")
        raise
    except json.JSONDecodeError as e:
        logger.error(f"Erro ao decodificar JSON: {e}")
        return {}
    except UnicodeDecodeError as e:
        logger.error(f"Erro de codificação ao ler arquivo: {e}")
        # Tenta com outra codificação
        try:
            with open(arquivo, 'r', encoding='utf-8') as f:
                dados = json.load(f)
                logger.info(f"Arquivo JSON carregado com sucesso usando UTF-8: {arquivo}")
                return dados
        except Exception as e2:
            logger.error(f"Falha ao tentar codificação alternativa: {e2}")
            return {}
    except Exception as e:
        logger.error(f"Erro inesperado ao carregar JSON: {e}")
        return {}

# Função para extrair as páginas do arquivo "Layout"
def extrair_paginas(layout: dict) -> str:
    """Extrai e organiza informações de páginas em formato Markdown."""
    markdown_output = [""]
    for section in layout.get('sections', []):
        page_name = section.get('displayName', 'Sem Nome')
        markdown_output.append(f"{page_name}\n-----------\n")
    return "\n".join(markdown_output)

# Função para extrair os visuais do arquivo "Layout"
def extrair_visuais(layout: dict) -> str:
    """Extrai e organiza informações de visuais em cada página em formato Markdown."""
    markdown_output = [""]
    for section in layout.get('sections', []):
        page_name = section.get('displayName', 'Sem Nome')
        for container in section.get("visualContainers", []):
            config_data = json.loads(container.get("config", "{}"))
            visual_type = config_data.get("singleVisual", {}).get("visualType")
            position = next(iter(config_data.get("layouts", [])), {}).get("position", {})
            query_refs = [item.get("queryRef") for items in config_data.get("singleVisual", {}).get("projections", {}).values()
                          for item in items if item.get("queryRef")]
            
            markdown_output.append(
                f"Página: {page_name}\n"
                f"X: {int(position.get('x', 0))}\n"
                f"Y: {int(position.get('y', 0))}\n"
                f"Altura: {int(position.get('height', 0))}\n"
                f"Largura: {int(position.get('width', 0))}\n"
                f"Tipo de visual: {visual_type}\n"
                f"Medidas utilizadas: {', '.join(query_refs) if query_refs else 'Não há medidas utilizadas no visual'}\n"
                "-----------\n"
            )
    return "\n".join(markdown_output)

def extrair_tabelas(model_data: dict) -> str:
    """Extrai e organiza informações de tabelas em formato Markdown."""
    markdown_output = [""]
    for table in model_data.get('model', {}).get('tables', []):
        table_name = table.get("name", "")
        if table_name.startswith("DateTableTemplate") or table_name.startswith("LocalDateTable"):
            continue
        for column in table.get('columns', []):
            column_name = column.get("name", "")
            data_type = column.get('dataType', "")
            is_calculated = 'Sim' if column.get('type', "") in ['calculatedTableColumn', 'calculated'] else 'Não'
            markdown_output.append(
                f"Tabela: {table_name}\n"
                f"Coluna: {column_name}\n"
                f"Tipo de dados: {data_type}\n"
                f"Coluna calculada?: {is_calculated}\n"
                "-----------\n"
            )
    return "\n".join(markdown_output)

def extrair_medidas(model_data: dict) -> str:
    """Extrai e organiza informações de medidas em formato Markdown."""
    markdown_output = [""]
    processed_measures = set()
    for table in model_data.get('model', {}).get('tables', []):
        table_name = table.get("name", "")
        for measure in table.get('measures', []):
            measure_name = measure.get('name', '')
            measure_expression = measure.get('expression', '')
            if (table_name, measure_name) in processed_measures:
                continue
            processed_measures.add((table_name, measure_name))
            if isinstance(measure_expression, list):
                measure_expression = ' '.join(filter(lambda x: x.strip(), measure_expression))
            markdown_output.append(
                f"Tabela: {table_name}\n"
                f"Medida: {measure_name}\n"
                f"Expressão: {measure_expression}\n"
                "-----------\n"
            )
    return "\n".join(markdown_output)

def extrair_fontes(model_data: dict) -> str:
    """Extrai e organiza informações sobre fontes de dados em formato Markdown."""
    markdown_output = [""]
    for table in model_data.get('model', {}).get('tables', []):
        table_name = table.get("name", "")
        if table_name.startswith("DateTableTemplate") or table_name.startswith("LocalDateTable"):
            continue
        for partition in table.get('partitions', []):
            partition_mode = partition.get('mode')
            source = partition.get('source', {})
            font_type = source.get('type')
            font_expression = source.get('expression')
            if isinstance(font_expression, list):
                font_expression = ' '.join(filter(lambda x: x.strip(), font_expression))
            markdown_output.append(
                f"Tabela: {table_name}\n"
                f"Modo de importação: {partition_mode}\n"
                f"Tipo de importação: {font_type}\n"
                f"Fonte: {font_expression}\n"
                "-----------\n"
            )
    return "\n".join(markdown_output)

def extrair_relacionamentos(model_data: dict) -> str:
    """Extrai e organiza informações de relacionamentos em formato Markdown."""
    markdown_output = [""]
    for relation in model_data.get('model', {}).get('relationships', []):
        from_table = relation.get('fromTable')
        to_table = relation.get('toTable')
        from_column = relation.get('fromColumn', '')
        to_column = relation.get('toColumn', '')
        if from_table.startswith("DateTableTemplate") or from_table.startswith("LocalDateTable") or \
           to_table.startswith("DateTableTemplate") or to_table.startswith("LocalDateTable"):
            continue
        markdown_output.append(
            f"Da tabela: {from_table}\n"
            f"Para tabela: {to_table}\n"
            f"Da coluna: {from_column}\n"
            f"Para coluna: {to_column}\n"
            "-----------\n"
        )
    return "\n".join(markdown_output)

"""
Segunda parte do código: 
1. Exportação para o Word das informações em Markdown
"""

def salvar_versao(salvar_path):
    """
    Verifica se o arquivo já existe e, se sim, cria um nome de arquivo com uma versão incremental.
    
    Parâmetros:
        salvar_path (str): Caminho completo do arquivo que deseja salvar.
        
    Retorna:
        str: Caminho final do arquivo com uma versão incrementada, se necessário.
    """

    # Se o arquivo não existir, retorna o caminho original
    if not os.path.exists(salvar_path):
        return salvar_path
    
    # Se já existir, adiciona a versão incremental
    base, ext = os.path.splitext(salvar_path)
    versao = 2
    
    # Incrementa a versão até encontrar um nome de arquivo disponível
    while os.path.exists(f"{base}_versão_{versao:02}{ext}"):
        versao += 1
    
    # Retorna o novo caminho do arquivo com a versão adicionada
    return f"{base}_versão_{versao:02}{ext}"

# Função para gerar o documento com conteúdo Markdown
def gerar_documento(nome_BI: str, extracoes: Dict[str, str], modelo_path: str, salvar_path: str) -> str:
    """
    Gera o documento Word com as descrições em formato Markdown nos locais apropriados.
    """
    # Carrega o modelo de documento do Word
    document = Document(modelo_path)

    # Preenche informações básicas do documento
    for para in document.paragraphs:
        if "Data da documentação:" in para.text:
            para.add_run(f" {datetime.now().strftime('%d/%m/%Y')}")
        elif "Nome do Relatório:" in para.text:
            para.add_run(f" {nome_BI}")

    # Insere cada seção Markdown no local correto do documento
    for titulo, conteudo_markdown in extracoes.items():
        for para in document.paragraphs:
            if para.text.strip() == titulo.capitalize():
                # Insere o conteúdo Markdown logo abaixo do parágrafo do título
                paragrafo_conteudo = document.add_paragraph(conteudo_markdown)
                para._element.addnext(paragrafo_conteudo._element)
                break  

    # Gera e salva o documento no caminho final, com controle de versão se necessário
    caminho_final = salvar_versao(salvar_path)
    document.save(caminho_final)
    print(f'Documentação gerada com sucesso em: {caminho_final}')
    return caminho_final

# Função principal para execução do processo
def main(arquivo_pbit: str, modelo_word: str, diretorio_saida: str) -> Optional[str]:
    """
    Função principal que coordena o processo de documentação.
    Extrai dados do arquivo Power BI e gera a documentação em Word.
    
    Args:
        arquivo_pbit (str): Caminho completo para o arquivo .pbit
        modelo_word (str): Caminho completo para o modelo Word
        diretorio_saida (str): Diretório onde será salvo o documento gerado
        
    Returns:
        Optional[str]: Caminho do arquivo gerado em caso de sucesso, None em caso de erro
    """
    try:
        logger.info("Iniciando processo de documentação")
        
        # Configurando caminhos
        nome_BI = Path(arquivo_pbit).stem
        caminho_BI = str(Path(arquivo_pbit).parent)
        arquivo_zip = os.path.join(caminho_BI, f'{nome_BI}.zip')
        
        # Criar diretório output se não existir
        output_dir = os.path.join(diretorio_saida, 'output')
        os.makedirs(output_dir, exist_ok=True)
        
        # Novo caminho do arquivo de saída
        salvar_path = os.path.join(output_dir, f'{nome_BI}_documentado.docx')
        
        logger.info(f"Processando relatório: {nome_BI}")
        
        # Verifica existência dos diretórios necessários
        for path in [caminho_BI, Path(modelo_word).parent, diretorio_saida]:
            if not os.path.exists(path):
                raise FileNotFoundError(f"Diretório não encontrado: {path}")
        
        # Verifica existência do modelo Word
        if not os.path.exists(modelo_word):
            raise FileNotFoundError(f"Modelo Word não encontrado: {modelo_word}")
        
        # Verifica e renomeia o arquivo para .zip, se necessário
        if not verificar_ou_renomear_arquivo(arquivo_pbit, arquivo_zip):
            logger.error("Falha ao verificar ou renomear o arquivo.")
            return None
        
        # Extrai arquivos do ZIP
        if not extrair_arquivos_zip(arquivo_zip, caminho_BI, ['Report/Layout', 'DataModelSchema']):
            logger.error("Falha ao extrair arquivos do ZIP.")
            return None
        
        # Carrega os dados JSON
        try:
            layout_data = carregar_dados_json(os.path.join(caminho_BI, 'Report/Layout'))
            model_data = carregar_dados_json(os.path.join(caminho_BI, 'DataModelSchema'))
        except FileNotFoundError as e:
            logger.error(f"Arquivo JSON não encontrado: {e}")
            return None
        
        # Reverte o arquivo ZIP para o formato original .pbit
        try:
            os.rename(arquivo_zip, arquivo_pbit)
            logger.info("Arquivo ZIP revertido para .pbit com sucesso")
        except Exception as e:
            logger.error(f"Erro ao reverter arquivo ZIP para .pbit: {e}")
            return None
        
        # Dicionário de extrações em Markdown
        extracoes = {
            "Páginas": extrair_paginas(layout_data),
            "Tabelas": extrair_tabelas(model_data),
            "Medidas": extrair_medidas(model_data),
            "Visuais": extrair_visuais(layout_data),
            "Fontes": extrair_fontes(model_data),
            "Relacionamentos": extrair_relacionamentos(model_data)
        }
        
        # Gera o documento final
        try:
            caminho_final = gerar_documento(nome_BI, extracoes, modelo_word, salvar_path)
            logger.info("Documentação gerada com sucesso!")
            return caminho_final
        except Exception as e:
            logger.error(f"Erro ao gerar documento: {e}")
            return None
            
    except Exception as e:
        logger.error(f"Erro inesperado durante a execução: {e}")
        return None
